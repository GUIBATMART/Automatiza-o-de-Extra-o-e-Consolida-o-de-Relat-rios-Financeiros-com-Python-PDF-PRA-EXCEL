# Projeto 6 - Automatizando a Geração de Relatórios Financeiros com Word, PDF, Excel e Python
# Gera Documento Word

# Importa a classe Document para criar e manipular documentos Word
from docx import Document

# Importa Pt para definir o tamanho da fonte dos textos
from docx.shared import Pt

# Importa nsdecls para lidar com namespaces XML no Word
from docx.oxml.ns import nsdecls

# Importa parse_xml para manipular elementos XML dentro do documento
from docx.oxml import parse_xml

# Cria um novo documento Word
doc = Document()

# Adiciona um título de nível 1 ao documento
doc.add_heading('Relatório Financeiro', level=1)

# Adiciona a primeira tabela com 11 linhas e 6 colunas ao documento
table1 = doc.add_table(rows=11, cols=6)

# Acessa a tabela no formato XML bruto
tbl1 = table1._tbl

# Acessa as propriedades da tabela
tbl1_pr = tbl1.tblPr

# Define as bordas da tabela usando XML para personalização
tbl1_borders = parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>' % nsdecls('w'))

# Aplica as bordas definidas à tabela
tbl1_pr.append(tbl1_borders)

# Define os cabeçalhos das colunas da primeira tabela
header_cells1 = table1.rows[0].cells
header_cells1[0].text = 'Data'
header_cells1[1].text = 'Descrição'
header_cells1[2].text = 'Receitas'
header_cells1[3].text = 'Despesas'
header_cells1[4].text = 'Saldo Anterior'
header_cells1[5].text = 'Saldo Atual'

# Ajusta o tamanho da fonte dos cabeçalhos da primeira tabela para 10 pontos
for cell in header_cells1:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)

# Define os dados a serem preenchidos nas linhas da primeira tabela
data1 = [
    ['01/08/2025', 'Venda Produto A', '5000', '', '10000', '15000'],
    ['02/08/2025', 'Compra Insumos', '', '2000', '15000', '13000'],
    ['03/08/2025', 'Venda Produto B', '3000', '', '13000', '16000'],
    ['04/08/2025', 'Despesas Operacionais', '', '1500', '16000', '14500'],
    ['05/08/2025', 'Venda Produto C', '7000', '', '14500', '21500'],
    ['06/08/2025', 'Despesas de Marketing', '', '2500', '21500', '19000'],
    ['07/08/2025', 'Venda Produto D', '4000', '', '19000', '23000'],
    ['08/08/2025', 'Compra de Equipamentos', '', '6000', '23000', '17000'],
    ['09/08/2025', 'Venda Produto E', '3500', '', '17000', '20500'],
    ['10/08/2024', 'Despesas de Transporte', '', '1200', '20500', '19300']
]

# Preenche as linhas da primeira tabela com os dados
for i, row_data in enumerate(data1, start=1):
    row_cells1 = table1.rows[i].cells
    for j, cell_data in enumerate(row_data):
        row_cells1[j].text = cell_data

# Adiciona uma quebra de página após a primeira tabela
doc.add_page_break()

# Adiciona a segunda tabela com 11 linhas e 6 colunas ao documento
table2 = doc.add_table(rows=11, cols=6)

# Acessa a tabela no formato XML bruto
tbl2 = table2._tbl

# Acessa as propriedades da tabela
tbl2_pr = tbl2.tblPr

# Define as bordas da segunda tabela usando XML para personalização
tbl2_borders = parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>' % nsdecls('w'))

# Aplica as bordas definidas à segunda tabela
tbl2_pr.append(tbl2_borders)

# Define os cabeçalhos das colunas da segunda tabela
header_cells2 = table2.rows[0].cells
header_cells2[0].text = 'Data'
header_cells2[1].text = 'Descrição'
header_cells2[2].text = 'Receitas'
header_cells2[3].text = 'Despesas'
header_cells2[4].text = 'Saldo Anterior'
header_cells2[5].text = 'Saldo Atual'

# Ajusta o tamanho da fonte dos cabeçalhos da segunda tabela para 10 pontos
for cell in header_cells2:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)

# Define os dados a serem preenchidos nas linhas da segunda tabela
data2 = [
    ['11/08/2025', 'Venda Produto F', '8000', '', '19300', '27300'],
    ['12/08/2025', 'Compra de Matéria-Prima', '', '4000', '27300', '23300'],
    ['13/08/2025', 'Venda Produto G', '4500', '', '23300', '27800'],
    ['14/08/2025', 'Despesas de Manutenção', '', '1700', '27800', '26100'],
    ['15/08/2025', 'Venda Produto H', '6000', '', '26100', '32100'],
    ['16/08/2025', 'Despesas Diversas', '', '2300', '32100', '29800'],
    ['17/08/2025', 'Venda Produto I', '7500', '', '29800', '37300'],
    ['18/08/2025', 'Compra de Ferramentas', '', '3200', '37300', '34100'],
    ['19/08/2025', 'Venda Produto J', '5000', '', '34100', '39100'],
    ['20/08/2025', 'Despesas Gerais', '', '1800', '39100', '37300']
]

# Preenche as linhas da segunda tabela com os dados
for i, row_data in enumerate(data2, start=1):
    row_cells2 = table2.rows[i].cells
    for j, cell_data in enumerate(row_data):
        row_cells2[j].text = cell_data

# Define o caminho para salvar o documento
file_path = "documento3.docx"

# Salva o documento no caminho especificado
doc.save(file_path)

# Retorna o caminho do arquivo salvo
file_path
