# Projeto 6 - Automatizando a Geração de Relatórios Financeiros com Word, PDF, Excel e Python
# Gera Documento Word

# Importa a classe Document para criar e manipular documentos Word
from docx import Document

# Importa Pt para definir o tamanho da fonte dos textos
from docx.shared import Pt

# Importa nsdecls para lidar com namespaces XML no Word
from docx.oxml.ns import nsdecls

# Importa parse_xml para manipular elementos XML dentro do documento
from docx.oxml import parse_xml

# Cria um novo documento Word
doc = Document()

# Adiciona um título de nível 1 ao documento
doc.add_heading('Relatório Financeiro', level=1)

# Adiciona uma tabela com 11 linhas e 6 colunas ao documento
table = doc.add_table(rows=11, cols=6)

# Acessa a tabela no formato XML bruto
tbl = table._tbl  

# Acessa as propriedades da tabela
tbl_pr = tbl.tblPr  

# Define as bordas da tabela usando XML para personalização
tbl_borders = parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>' % nsdecls('w'))

# Aplica as bordas definidas à tabela
tbl_pr.append(tbl_borders)

# Define os cabeçalhos das colunas da tabela
header_cells = table.rows[0].cells
header_cells[0].text = 'Data'
header_cells[1].text = 'Descrição'
header_cells[2].text = 'Receitas'
header_cells[3].text = 'Despesas'
header_cells[4].text = 'Saldo Anterior'
header_cells[5].text = 'Saldo Atual'

# Ajusta o tamanho da fonte dos cabeçalhos da tabela para 10 pontos
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)

# Define os dados a serem preenchidos nas linhas da tabela
data = [
    ['01/08/2025', 'Venda Produto A', '5000', '', '10000', '15000'],
    ['02/08/2025', 'Compra Insumos', '', '2000', '15000', '13000'],
    ['03/08/2025', 'Venda Produto B', '3000', '', '13000', '16000'],
    ['04/08/2025', 'Despesas Operacionais', '', '1500', '16000', '14500'],
    ['05/08/2025', 'Venda Produto C', '7000', '', '14500', '21500'],
    ['06/08/2025', 'Despesas de Marketing', '', '2500', '21500', '19000'],
    ['07/08/2025', 'Venda Produto D', '4000', '', '19000', '23000'],
    ['08/08/2025', 'Compra de Equipamentos', '', '6000', '23000', '17000'],
    ['09/08/2025', 'Venda Produto E', '3500', '', '17000', '20500'],
    ['10/08/2024', 'Despesas de Transporte', '', '1200', '20500', '19300']
]

# Preenche as linhas da tabela com os dados
for i, row_data in enumerate(data, start=1):
    row_cells = table.rows[i].cells
    for j, cell_data in enumerate(row_data):
        row_cells[j].text = cell_data

# Define o caminho para salvar o documento
file_path = "documento2.docx"

# Salva o documento no caminho especificado
doc.save(file_path)

# Retorna o caminho do arquivo salvo
file_path
